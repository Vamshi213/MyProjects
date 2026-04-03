"""
Resume parser — extracts structured data from PDF, DOCX, or plain text.
Returns raw text + structured fields (skills, years_exp, titles, etc.)
"""
import re
import os
from typing import Dict, List, Optional, Tuple

from .skills_db import ALL_SKILLS, STRONG_VERBS, WEAK_VERBS


# ── Text extraction ───────────────────────────────────────────────────────────

def extract_text(file_path: str, file_type: str) -> str:
    """Extract plain text from uploaded resume file."""
    ext = file_type.lower().strip(".")
    try:
        if ext == "pdf":
            return _extract_pdf(file_path)
        elif ext in ("docx",):
            return _extract_docx(file_path)
        elif ext in ("doc",):
            return _extract_doc_fallback(file_path)
        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception as e:
        return ""


def _extract_pdf(path: str) -> str:
    try:
        import PyPDF2
        text_parts = []
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)
    except Exception:
        return ""


def _extract_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_doc_fallback(path: str) -> str:
    """Best-effort plain text extraction for legacy .doc files."""
    try:
        with open(path, "rb") as f:
            raw = f.read()
        text = raw.decode("utf-8", errors="ignore")
        # Strip non-printable characters
        text = re.sub(r"[^\x20-\x7E\n\t]", " ", text)
        return text
    except Exception:
        return ""


# ── Structured parsing ────────────────────────────────────────────────────────

def parse_resume(file_path: str, file_type: str) -> Dict:
    """
    Parse resume and return structured dict:
    {
        text, skills, years_experience, job_titles, education,
        certifications, has_metrics, has_links, bullet_count,
        weak_verbs_found, strong_verbs_found, word_count
    }
    """
    text = extract_text(file_path, file_type)
    if not text.strip():
        return {"error": "Could not extract text from resume", "text": ""}

    text_lower = text.lower()
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    return {
        "text": text,
        "text_lower": text_lower,
        "skills": _extract_skills(text_lower),
        "years_experience": _estimate_years(text),
        "job_titles": _extract_titles(text),
        "education": _extract_education(text_lower),
        "certifications": _extract_certifications(text_lower),
        "has_metrics": _has_metrics(text),
        "metric_count": _count_metrics(text),
        "has_links": bool(re.search(r"github\.com|linkedin\.com|portfolio|website", text_lower)),
        "bullet_count": _count_bullets(lines),
        "weak_verbs_found": _find_weak_verbs(text_lower),
        "strong_verbs_found": _find_strong_verbs(text_lower),
        "word_count": len(text.split()),
        "line_count": len(lines),
    }


def _extract_skills(text_lower: str) -> List[str]:
    found = []
    for skill in ALL_SKILLS:
        pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, text_lower):
            found.append(skill)
    return found


def _estimate_years(text: str) -> int:
    """Estimate total years of experience from date ranges."""
    # Look for year ranges like 2019 - 2023 or 2019–Present
    patterns = [
        r"(20\d\d)\s*[-–—]\s*(20\d\d|present|current)",
        r"(19\d\d)\s*[-–—]\s*(20\d\d)",
    ]
    years_total = 0
    current_year = 2026
    for pat in patterns:
        matches = re.findall(pat, text, re.IGNORECASE)
        for start, end in matches:
            try:
                s = int(start)
                e = current_year if end.lower() in ("present", "current") else int(end)
                years_total += max(0, e - s)
            except ValueError:
                pass
    return min(years_total, 30)  # cap sanity


def _extract_titles(text: str) -> List[str]:
    title_patterns = [
        r"(senior|staff|principal|lead|junior|associate)?\s*(software|backend|frontend|full[ -]?stack|"
        r"platform|infrastructure|devops|cloud|data|ml|machine learning|ai|mobile|ios|android|"
        r"security|site reliability)\s*(engineer|developer|architect|scientist|analyst)?",
        r"(engineering|product|project|program|technical|general)\s*(manager|director|lead|head|vp|"
        r"vice president)",
        r"(cto|ceo|vp|head of engineering|head of product|principal engineer|staff engineer|"
        r"distinguished engineer|fellow)",
        r"(ux|ui|product|graphic)\s*(designer|researcher|strategist)",
    ]
    found = set()
    text_lower = text.lower()
    for pat in title_patterns:
        for m in re.finditer(pat, text_lower):
            title = m.group(0).strip()
            if len(title) > 4:
                found.add(title)
    return list(found)[:10]


def _extract_education(text_lower: str) -> Dict:
    degrees = {
        "phd": ["phd", "ph.d", "doctorate", "doctor of philosophy"],
        "masters": ["master", "m.s.", "msc", "mba", "m.eng"],
        "bachelors": ["bachelor", "b.s.", "bsc", "b.e.", "b.tech", "undergraduate"],
        "bootcamp": ["bootcamp", "boot camp", "coding school"],
    }
    found_degrees = []
    for level, keywords in degrees.items():
        if any(k in text_lower for k in keywords):
            found_degrees.append(level)

    cs_fields = ["computer science", "software engineering", "information technology",
                 "electrical engineering", "mathematics", "statistics", "data science"]
    has_cs_degree = any(f in text_lower for f in cs_fields)

    return {
        "degrees": found_degrees,
        "has_cs_degree": has_cs_degree,
        "highest": found_degrees[0] if found_degrees else "unknown",
    }


def _extract_certifications(text_lower: str) -> List[str]:
    cert_keywords = [
        "aws certified", "google cloud", "azure certified", "cka", "ckad",
        "cissp", "ceh", "oscp", "terraform associate", "databricks",
        "pmp", "scrum master", "csm", "comptia",
    ]
    return [c for c in cert_keywords if c in text_lower]


def _has_metrics(text: str) -> bool:
    return bool(re.search(r"\d+[%xX]|\$[\d,]+|\d+[kmb]\b|\d+\s*(million|billion|thousand)", text, re.IGNORECASE))


def _count_metrics(text: str) -> int:
    patterns = [
        r"\d+%",
        r"\$[\d,]+[kmb]?",
        r"\d+[xX]\b",
        r"\d+\s*(million|billion|k|m)\b",
        r"\d{1,3}(?:,\d{3})+",  # numbers with commas like 1,000,000
    ]
    count = 0
    for pat in patterns:
        count += len(re.findall(pat, text, re.IGNORECASE))
    return count


def _count_bullets(lines: List[str]) -> int:
    bullet_starts = ["•", "-", "·", "◦", "▸", "→", "*"]
    return sum(1 for l in lines if any(l.startswith(b) for b in bullet_starts)
               or re.match(r"^\s*[-•·◦▸→*]\s", l))


def _find_weak_verbs(text_lower: str) -> List[str]:
    found = []
    for verb in WEAK_VERBS:
        if re.search(r"\b" + re.escape(verb) + r"\b", text_lower):
            found.append(verb)
    return found


def _find_strong_verbs(text_lower: str) -> List[str]:
    found = []
    for verb in STRONG_VERBS:
        if re.search(r"\b" + re.escape(verb) + r"\b", text_lower):
            found.append(verb)
    return found
